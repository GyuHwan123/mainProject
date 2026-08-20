from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "finance-evaluation-manual.md"
OUTPUT = ROOT / "docs" / "영수증_LLM_비교평가_사용매뉴얼.docx"

NAVY = "17365D"
BLUE = "2E74B5"
PALE_BLUE = "E8EEF5"
PALE_YELLOW = "FFF4CC"
GRAY = "65758B"
LIGHT_GRAY = "F4F6F9"
RED = "9B1C1C"


def set_font(run, name="맑은 고딕", size=10.5, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, fill, border=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border:
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:color"), border)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("페이지 ")
    set_font(run, size=8.5, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def add_inline(paragraph, text, base_size=10.5, color=None):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, name="Consolas", size=9.5, color=NAVY)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=base_size, bold=True, color=color)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=base_size, color=color)


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    add_inline(p, text)


def add_list(doc, text, ordered=False):
    p = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    add_inline(p, text)


def add_code(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F5F7")
    tc_pr.append(shd)
    margins = OxmlElement("w:tcMar")
    for edge, value in (("top", "100"), ("bottom", "100"), ("start", "140"), ("end", "140")):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), value)
        el.set(qn("w:type"), "dxa")
        margins.append(el)
    tc_pr.append(margins)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(lines):
        if index:
            p.add_run().add_break()
        run = p.add_run(line)
        set_font(run, name="Consolas", size=8.5, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)

    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    specs = {
        "Heading 1": (16, BLUE, 18, 8),
        "Heading 2": (13, BLUE, 14, 6),
        "Heading 3": (11.5, NAVY, 10, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "맑은 고딕"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "맑은 고딕"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(10.5)


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    hp = section.header.paragraphs[0]
    hp.text = "DOCUNEX  |  FINANCE MODEL LAB"
    set_font(hp.runs[0], size=8.5, bold=True, color=GRAY)
    add_page_number(section.footer.paragraphs[0])

    # Editorial-style first page header.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FINANCE MODEL LAB")
    set_font(r, size=11, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("영수증 LLM 비교 평가 페이지")
    set_font(r, size=27, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    r = subtitle.add_run("사용 및 로컬 Ollama 모델 적용 매뉴얼")
    set_font(r, size=15, color=BLUE)

    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.paragraph_format.space_after = Pt(28)
    r = lead.add_run("OCR 결과 · 모델별 필드 정확도 · Excel 변환 결과를 동일 조건으로 비교하는 운영 안내서")
    set_font(r, size=10.5, color=GRAY)

    note = doc.add_paragraph()
    note.paragraph_format.left_indent = Inches(0.35)
    note.paragraph_format.right_indent = Inches(0.35)
    note.paragraph_format.space_before = Pt(14)
    note.paragraph_format.space_after = Pt(10)
    note.paragraph_format.line_spacing = 1.25
    shade_paragraph(note, PALE_YELLOW, "D6A700")
    add_inline(note, "핵심: PPUF가 아니라 GGUF입니다. 어댑터 파일의 확장자만 바꾸는 것이 아니라, llama.cpp 변환기로 실제 GGUF 형식으로 변환한 후 Ollama에 등록해야 합니다.", base_size=11, color=NAVY)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(55)
    r = meta.add_run("대상: 개발자 및 관리자  |  환경: Windows PowerShell + Ollama")
    set_font(r, size=9.5, color=GRAY)

    doc.add_page_break()

    # Skip source title and opening duplicate note; retain substantive content.
    index = 1
    in_code = False
    code_lines = []
    skipped_intro_note = False
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        index += 1
        if stripped.startswith("```"):
            if in_code:
                add_code(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(raw)
            continue
        if not stripped:
            continue
        if stripped.startswith(">"):
            if not skipped_intro_note:
                skipped_intro_note = True
                continue
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(8)
            shade_paragraph(p, PALE_YELLOW)
            add_inline(p, stripped.lstrip("> "), color=NAVY)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            heading = doc.add_heading(stripped[3:], level=1)
            if re.match(r"\d+\.", stripped[3:]):
                heading.paragraph_format.page_break_before = False
        elif re.match(r"^\d+\. ", stripped):
            add_list(doc, re.sub(r"^\d+\. ", "", stripped), ordered=True)
        elif stripped.startswith("- "):
            add_list(doc, stripped[2:], ordered=False)
        else:
            add_body(doc, stripped)

    # Sources as a compact closing block.
    for p in doc.paragraphs:
        if p.text.strip() == "참고 자료":
            shade_paragraph(p, PALE_BLUE)

    props = doc.core_properties
    props.title = "영수증 LLM 비교 평가 페이지 사용 매뉴얼"
    props.subject = "로컬 Ollama GGUF 어댑터 등록 및 모델 평가 절차"
    props.author = "Docunex"
    props.keywords = "Ollama, GGUF, QLoRA, LoRA, OCR, LLM 평가"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
