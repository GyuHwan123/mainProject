from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(r"C:\Users\2Class_08\Desktop\main\mainProject\ocr-web-app\reports\gemma2_2b_4000개_학습_실험_1차_보고서.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK = "1F4D78"
INK = "24364B"
MUTED = "6B7785"
LIGHT = "F2F4F7"
PALE = "E8EEF5"
GREEN = "1D6F42"
RED = "9B1C1C"
GOLD = "7A5A00"
FONT = "Malgun Gothic"


def set_font(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        shade(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(str(text)), size=font_size, bold=True, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(str(value)), size=font_size, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_p(doc, text="", bold_lead=None, color=INK, after=6, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        set_font(p.add_run(bold_lead), size=size, bold=True, color=color)
        set_font(p.add_run(text[len(bold_lead):]), size=size, color=color)
    else:
        set_font(p.add_run(text), size=size, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    set_font(p.add_run(text), size=10.3)
    return p


def add_callout(doc, label, text, fill=PALE, label_color=DARK):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(f"{label}  "), size=10.5, bold=True, color=label_color)
    set_font(p.add_run(text), size=10.5, color=INK)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_font(run, size=9, color=MUTED)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = section.bottom_margin = Inches(1)
section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in (
    ("Title", 25, INK, 0, 6),
    ("Subtitle", 13, MUTED, 0, 18),
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, DARK, 8, 4),
):
    style = styles[name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = name != "Subtitle"
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("GEMMA2:2B 영수증 문서화 실험 | 1차"), size=8.5, color=MUTED)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(footer.add_run("실험 보고서  |  "), size=9, color=MUTED)
add_page_number(footer)

# Cover / opening block: editorial report pattern with restrained business styling.
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(80)
p.paragraph_format.space_after = Pt(12)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("1차 실험 보고서"), size=11, bold=True, color=BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
set_font(p.add_run("Gemma2:2B 4,000개 학습 실험"), size=25, bold=True, color=INK)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(28)
set_font(p.add_run("한국어 영수증 OCR · QLoRA · 재무 Excel 자동화 비교 평가"), size=13, color=MUTED)

add_table(doc, ["항목", "내용"], [
    ["평가 일자", "2026-08-19"],
    ["평가 데이터", "한국어 영수증 17건 (receipt_kr.json)"],
    ["비교 모델", "gemma2:2b / gemma2-no-category / gemma2-yes-category"],
    ["평가 방식", "순수 모델 결과 + 규칙 보정·Excel 생성 후 최종 시스템 결과"],
    ["총 호출", "17건 × 3개 모델 = 51회"],
], [1900, 7460], font_size=9.7)

add_callout(doc, "핵심 결론", "현재 평가에서는 기본 gemma2:2b가 최종 정확도 36.76%로 1위였다. 다만 이 수치는 품목 중심·엄격 일치 평가에 크게 좌우되며, yes 모델이 목표로 한 카드사·승인번호·마스킹 카드번호·공급가액·부가세 능력은 현재 정답 데이터로 충분히 측정되지 않았다.")

doc.add_page_break()
doc.add_heading("1. 실험 목적과 시스템 구성", level=1)
add_p(doc, "본 실험의 목적은 기본 Gemma2:2B와 4,000개 학습 데이터로 만든 두 QLoRA 어댑터를 동일한 한국어 영수증 OCR 입력에 적용하고, 순수 모델 추출 성능과 실제 Excel 문서화 시스템 성능을 분리해 비교하는 것이다.")
doc.add_heading("1.1 비교 모델", level=2)
add_table(doc, ["모델", "역할", "주요 의도"], [
    ["gemma2:2b", "기준 모델", "별도 미세조정 없는 Ollama 기본 모델"],
    ["gemma2-no-category", "QLoRA 이전/대조 모델", "카테고리 정보 강조 전 학습 모델"],
    ["gemma2-yes-category", "QLoRA 신규 모델", "결제수단, 카드정보, 공급가액·부가세·합계 등 Excel 입력 핵심값 강화"],
], [2100, 2100, 5160])
doc.add_heading("1.2 실행 구조", level=2)
add_p(doc, "평소 개발 환경에서는 프런트엔드를 npm으로, 백엔드와 OCR 서비스를 각각 로컬 가상환경의 uvicorn(8000, 8001 포트)으로 실행했다. 이 경우 LLM 호출은 Docker가 아니라 Windows 로컬 Ollama(127.0.0.1:11434)를 사용한다.")
add_bullet(doc, "OCR: 로컬 OCR 서비스가 이미지에서 텍스트와 근거 영역을 추출")
add_bullet(doc, "LLM: 동일 OCR 텍스트를 세 모델에 각각 전달하여 JSON 생성")
add_bullet(doc, "최종 시스템: 날짜·금액 규칙 보정과 정규화 적용")
add_bullet(doc, "Excel: Python openpyxl 기반 finance_workbook_service가 4종 재무 양식 생성")
add_callout(doc, "환경 구분", "로컬 백엔드는 로컬 Ollama를 사용한다. Docker Compose 백엔드는 Docker 내부의 ollama:11434를 사용하며 두 Ollama의 모델 저장소는 서로 독립적이다.")

doc.add_heading("2. 평가 설계", level=1)
doc.add_heading("2.1 평가 데이터", level=2)
add_p(doc, "receipt_kr.json은 17건의 영수증에 대해 가게명, 구매일자, 구매물품, 총 물품 수량, 총 결제액, 결제방식, 할부 여부를 제공한다. 이미지 파일명이 JSON에 포함되어 있지 않아 0~16번 순서로 이미지와 정답을 연결했다.")
doc.add_heading("2.2 순수 모델과 최종 시스템", level=2)
add_table(doc, ["구분", "포함 단계", "해석"], [
    ["순수 모델", "OCR 텍스트 → LLM JSON", "미세조정 모델 자체의 구조화 능력"],
    ["최종 시스템", "LLM JSON → 규칙 보정 → 정규화 → Excel", "실제 서비스에서 사용자가 받는 값"],
], [1800, 3600, 3960])
doc.add_heading("2.3 현재 점수 구성", level=2)
add_p(doc, "각 모델은 총 204개 필드로 평가되었다. 일반 필드 84개와 30개 품목의 4개 하위 필드(품목명·수량·단가·금액) 120개로 구성된다. 품목 평가가 전체의 58.8%를 차지한다.")
add_callout(doc, "해석 주의", "현재 정확도는 공백·특수문자·OCR 유사문자 차이도 오답으로 처리하는 엄격 일치 점수다. 품목은 배열 순서대로 비교되어 한 행이 추가되거나 누락되면 뒤 행까지 연쇄 오답이 될 수 있다.", fill="FFF8E8", label_color=GOLD)

doc.add_heading("3. 17건 전체 결과", level=1)
add_table(doc, ["모델", "순수", "최종", "최종 정답", "완전 정답", "평균 응답"], [
    ["gemma2:2b", "37.25%", "36.76%", "75/204", "0/17", "46.42초"],
    ["gemma2-no-category", "19.61%", "20.10%", "41/204", "0/17", "33.39초"],
    ["gemma2-yes-category", "14.22%", "16.18%", "33/204", "0/17", "33.95초"],
], [2350, 1200, 1200, 1500, 1500, 1610])
add_p(doc, "기본 gemma2:2b가 최종 시스템 정확도 36.76%로 가장 높았다. no-category는 기본 모델보다 16.66%p, yes-category는 20.58%p 낮았다. 반면 학습 모델들의 평균 응답시간은 기본 모델보다 약 12~13초 짧았다.")
doc.add_heading("3.1 최종 시스템 필드별 정확도", level=2)
add_table(doc, ["필드", "기본", "no-category", "yes-category"], [
    ["결제일", "88.2%", "88.2%", "70.6%"],
    ["합계금액", "76.5%", "76.5%", "76.5%"],
    ["상호명", "41.2%", "17.6%", "23.5%"],
    ["결제방식", "5.9%", "0%", "5.9%"],
    ["카테고리", "0%", "0%", "0%"],
    ["품목명", "10.0%", "6.7%", "3.3%"],
    ["품목 수량", "60.0%", "13.3%", "0%"],
    ["품목 단가", "20.0%", "10.0%", "3.3%"],
    ["품목 금액", "40.0%", "3.3%", "3.3%"],
], [3000, 2120, 2120, 2120])
doc.add_heading("3.2 구조적 오류", level=2)
add_bullet(doc, "no-category는 17건 중 15건에서 items가 없거나 빈 배열이었다.")
add_bullet(doc, "yes-category는 17건 모두에서 items가 없거나 빈 배열이었다.")
add_bullet(doc, "yes-category는 17건 모두 document_type 대신 doc_type을 반환했다.")
add_bullet(doc, "세 모델 모두 LLM 호출과 Excel 파일 생성에는 17/17건 성공했다.")
add_callout(doc, "중요", "Excel 생성 성공은 파일이 열리고 예상 시트가 활성화되었다는 뜻이다. 행 수와 셀 값이 정답이라는 뜻은 아니다.", fill="FDECEC", label_color=RED)

doc.add_heading("4. 정확도가 낮은 원인", level=1)
doc.add_heading("4.1 평가 목적과 정답 데이터의 불일치", level=2)
add_p(doc, "yes 모델은 카드/현금 구분, 카드사, 마스킹 카드번호, 승인번호, 공급가액, 부가세, 합계금액을 잘 인식하도록 학습했지만 receipt_kr.json에는 카드사·승인번호·공급가액·부가세가 독립 정답 필드로 존재하지 않는다. 현재 직접 평가되는 yes 모델 핵심 필드는 결제방식 17개와 합계금액 17개, 총 34개로 전체 204개 중 약 16.7%에 불과하다.")
doc.add_heading("4.2 품목 중심 가중치", level=2)
add_p(doc, "품목 필드가 전체 점수의 58.8%를 차지한다. 품목을 전혀 출력하지 않으면 나머지 일반 필드를 모두 맞혀도 이론적 최고점은 약 41.2%다. 따라서 결제정보 특화 모델을 품목 중심 전체 점수로 평가하는 것은 불공정하다.")
doc.add_heading("4.3 엄격 일치와 순서 비교", level=2)
add_p(doc, "‘★LG 34UC79G. BKR(핸)’과 ‘LG 34UC79G.BKR(한)’처럼 의미상 동일하거나 OCR 오차로 설명 가능한 차이도 오답이다. ‘카드(3562**153*)’와 ‘CARD’도 결제수단 종류는 맞지만 전체 문자열이 다르므로 오답이다.")
doc.add_heading("4.4 학습·서비스 스키마 불일치", level=2)
add_p(doc, "현재 서비스는 document_type과 items 배열을 요구한다. yes 모델이 doc_type을 일관되게 반환하고 items를 전부 생략한 것은 단순 학습 부족보다는 학습 정답 형식이 서비스 프롬프트와 다르다는 강한 신호다. doc_type 문제와 items 누락은 별개의 문제이며 키 이름만 고친다고 품목이 자동으로 복구되지는 않는다.")
doc.add_heading("4.5 어댑터 배포 방식", level=2)
add_p(doc, "어댑터는 google/gemma-2-2b-it 기반으로 학습됐고, Ollama에서는 Q4_0 양자화된 gemma2:2b에 직접 결합했다. 베이스 체크포인트와 양자화 방식의 차이가 품질 저하를 일으킬 가능성이 있다. 정확한 원본 베이스에 어댑터를 병합한 뒤 GGUF로 변환한 모델과 비교할 필요가 있다.")

doc.add_heading("5. OCR 영향과 분리 평가", level=1)
add_p(doc, "OCR 성능 부족은 일부 오류에 영향을 주었을 가능성이 있다. 상품명의 글자 변형, 불필요한 단어를 품목으로 추가한 사례가 그 정황이다. 그러나 현재 평가 내보내기에는 모델에 전달된 전체 OCR 원문과 필드별 OCR confidence가 포함되지 않아 OCR의 기여도를 정량화할 수 없다.")
add_table(doc, ["비교 입력", "결과 해석"], [
    ["실제 OCR 텍스트", "현재 전체 파이프라인 성능"],
    ["사람이 전사한 깨끗한 텍스트", "OCR 영향을 제거한 LLM 추출 성능"],
    ["두 입력의 차이", "OCR로 인한 손실량"],
], [3200, 6160])
add_bullet(doc, "깨끗한 텍스트에서는 맞고 OCR에서 틀리면 OCR 문제다.")
add_bullet(doc, "두 입력에서 모두 틀리면 LLM·학습 또는 스키마 문제다.")
add_bullet(doc, "OCR에 값 자체가 없으면 모델이 복원할 근거가 없으며, 임의 복원은 환각이다.")
add_p(doc, "날짜와 합계금액 점수가 높은 이유도 LLM만의 성능으로 해석하면 안 된다. 최종 시스템은 정규식 기반 _receipt_hints()로 날짜와 금액을 보정한다.")

doc.add_heading("6. Excel 신뢰성 평가", level=1)
doc.add_heading("6.1 현재 확인된 범위", level=2)
add_p(doc, "현재 Excel 검증은 파일 생성, 4개 시트 존재, 예측 유형에 해당하는 시트 활성화를 확인한다. 실제 데이터 행 수, 셀 내용, 누락과 환각은 확인하지 않는다. 특히 구매요청서 행 수는 items 배열에 의존하므로 학습 모델의 품목 누락은 Excel 행 누락으로 직결될 수 있다.")
doc.add_heading("6.2 다음 평가에서 필요한 셀 지표", level=2)
add_table(doc, ["지표", "정의"], [
    ["대상 셀", "정답상 값이 들어가야 하는 동적 데이터 셀"],
    ["채움률", "채워진 필수 셀 ÷ 값이 필요한 정답 셀"],
    ["셀 정확도", "정확한 셀 ÷ 비교 대상 셀"],
    ["입력 정밀도", "정확한 입력 셀 ÷ 모델이 채운 셀"],
    ["입력 재현율", "정확한 입력 셀 ÷ 값이 필요한 정답 셀"],
    ["환각률", "정답은 null인데 불필요하게 채운 셀 비율"],
    ["행 정확도", "정답 품목 수와 Excel 데이터 행 수의 일치"],
], [2400, 6960])
add_p(doc, "정답도 null이고 셀도 빈 경우는 정상 공란이다. 필수 값인데 비어 있으면 누락이고, 정답이 null인데 값이 들어가면 환각이다. 정적 제목·결재란·빈 템플릿 시트는 평가 대상 셀에서 제외해야 한다.")
doc.add_heading("6.3 신뢰도 산정", level=2)
add_p(doc, "LLM이 스스로 출력한 확률을 그대로 신뢰해서는 안 된다. 필드 신뢰도는 OCR 근거, 모델 간 합의, 규칙 검산, 실제 검증 세트 정확도를 결합하고 검증 데이터로 보정해야 한다.")
add_bullet(doc, "OCR 근거: 값이 OCR 텍스트와 bbox에 존재하며 OCR confidence가 충분한가")
add_bullet(doc, "모델 합의: 세 모델이 동일한 값을 냈는가")
add_bullet(doc, "규칙 검산: 공급가액 + 부가세 = 합계, 수량 × 단가 = 품목 금액인가")
add_bullet(doc, "실측 신뢰도: 별도 검증 세트에서 해당 컬럼의 실제 정확도가 얼마인가")
add_callout(doc, "권장 출력", "문서별로 대상 셀·채운 셀·정답·오답·누락·환각 수와 컬럼별 정확도를 제공하고, 신뢰도는 높음/보통/낮음과 그 근거를 함께 표시한다.")

doc.add_heading("7. yes 모델 목표에 맞는 평가 스키마", level=1)
add_p(doc, "결제정보 특화 능력을 검증하려면 정답 데이터를 다음처럼 분리해야 한다. 원본에 값이 없으면 키를 생략하지 않고 null을 사용해야 환각도 측정할 수 있다.")
add_table(doc, ["영역", "권장 필드"], [
    ["기본", "merchant, transaction_date, transaction_time"],
    ["결제", "payment_method_type, card_company, masked_card_number, approval_number, installment"],
    ["금액", "supply_amount, tax_amount, total_amount, currency"],
    ["업무 분류", "document_type, expense_category"],
    ["품목", "items[].name, quantity, unit_price, supply_amount, tax_amount, total_amount"],
], [2200, 7160])
add_p(doc, "현재 서비스 프롬프트도 card_company, masked_card_number, approval_number, installment 등을 독립 키로 요구하지 않으며, Excel 양식에도 이 값들의 독립 열이 없다. 평가 데이터뿐 아니라 프롬프트, 저장 스키마, Excel 열 매핑까지 같은 구조로 맞춰야 실제 업무 성능을 검증할 수 있다.")

doc.add_heading("8. 추가 학습과 모델 크기 판단", level=1)
doc.add_heading("8.1 4,000개·1 epoch의 의미", level=2)
add_p(doc, "현재 오류는 학습량 부족보다 학습 목표와 서비스 스키마 불일치 신호가 더 강하다. 잘못된 키와 누락 구조를 가진 데이터로 epoch만 늘리면 잘못된 형식을 더 강하게 학습할 수 있다.")
add_bullet(doc, "스키마와 프롬프트를 먼저 통일한다.")
add_bullet(doc, "학습·검증·테스트 세트를 분리하고 17건은 학습에 포함하지 않는다.")
add_bullet(doc, "동일 데이터로 1·2·3 epoch 체크포인트를 저장해 validation loss와 필드 성능을 비교한다.")
add_bullet(doc, "JSON 파싱 성공률, 필수 키 존재율, items 생성률, 환각률을 함께 본다.")
doc.add_heading("8.2 Gemma 3 4B 가능성", level=2)
add_p(doc, "Gemma 3 4B는 더 높은 표현력과 지시 준수 능력으로 복수 품목과 한국어 OCR 노이즈 처리에서 개선될 가능성이 있다. 그러나 현재 Gemma2 어댑터를 재사용할 수 없고, 데이터 스키마가 틀리면 더 큰 모델도 같은 오류를 학습한다. 먼저 Gemma2 2B 파이프라인을 바로잡은 뒤 동일 조건으로 4B를 비교해야 모델 크기 효과를 분리할 수 있다.")

doc.add_heading("9. 권고 실행계획", level=1)
add_table(doc, ["단계", "작업", "완료 기준"], [
    ["1", "정답·서비스·Excel 스키마 통일", "모든 학습 샘플이 document_type과 전체 키를 일관되게 포함"],
    ["2", "결제정보 전용 정답 보강", "카드사·승인번호·마스킹 번호·공급가액·부가세를 null 포함 독립 필드화"],
    ["3", "평가 개선", "엄격/정규화 점수 분리, 품목 순서 독립 매칭, 목적별 점수 제공"],
    ["4", "OCR 영향 분리", "OCR 원문·confidence 저장 및 깨끗한 전사문 대비 성능 산출"],
    ["5", "Excel 셀 검증", "행·셀 단위 정답/오답/누락/환각과 컬럼별 신뢰도 산출"],
    ["6", "배포 방식 비교", "직접 어댑터 방식과 원본 베이스 병합→GGUF 방식 비교"],
    ["7", "재학습", "1·2·3 epoch 검증 곡선과 독립 테스트 결과 확보"],
    ["8", "4B 비교", "수정된 동일 데이터·프롬프트로 2B 대비 비용/성능 비교"],
], [900, 3300, 5160], font_size=8.7)

doc.add_heading("10. 최종 총평", level=1)
add_p(doc, "이번 1차 실험은 학습 모델이 기본 모델보다 낮은 전체 점수를 냈다는 사실뿐 아니라, 현재 평가가 yes 모델의 결제정보 특화 목표를 충분히 측정하지 못하고 있다는 점을 확인했다는 데 의미가 있다. Excel 파일은 17건 모두 생성됐지만, 이는 구조적 생성 성공일 뿐 셀 내용의 정확성을 보증하지 않는다.")
add_p(doc, "즉시 서비스 기본 모델을 교체하기보다는 gemma2:2b를 유지하고, 결제정보 중심 정답 스키마와 Excel 셀 단위 평가를 먼저 구축하는 것이 타당하다. 이후 동일 프롬프트·동일 스키마·정확한 베이스 병합 조건으로 Gemma2 2B를 재학습하고, 성능 개선이 확인된 다음 Gemma 3 4B 확장을 검토해야 한다.")
add_callout(doc, "의사결정", "현재 학습 모델은 연구 후보로 유지하되 운영 기본 모델로 승격하지 않는다. 다음 실험의 성공 기준은 전체 점수 하나가 아니라 결제정보 추출 정확도, 품목 성능, OCR 독립 성능, Excel 셀 정확도와 환각률을 각각 개선하는 것이다.")

doc.add_heading("부록 A. 실험 산출물", level=1)
add_table(doc, ["산출물", "설명"], [
    ["finance-model-evaluation-combined-17.json", "17개 영수증 × 3개 모델의 누적 원본 결과와 통계"],
    ["receipt_kr.json", "현재 평가에 사용한 한국어 영수증 구조화 정답"],
    ["재무 모델 평가 페이지", "세 모델의 순수·최종 결과 및 필드 비교 화면"],
    ["4종 Excel 양식", "경비지출결의서, 출장여비교통비정산서, 구매품의요청서, 복리후생비신청서"],
], [3300, 6060])

doc.core_properties.title = "Gemma2:2B 4,000개 학습 실험 1차 보고서"
doc.core_properties.subject = "한국어 영수증 OCR, QLoRA, 재무 Excel 자동화 비교 평가"
doc.core_properties.author = "PicToText 프로젝트"
doc.core_properties.keywords = "Gemma2, QLoRA, OCR, 영수증, Excel, 평가"
doc.save(OUT)
print(OUT)
