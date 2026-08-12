from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZipFile

from docx import Document

from app.schemas.ocr import OCRPage


def extract_docx_text(file_path: Path) -> str:
    """Extract native paragraph and table text from a DOCX file."""
    document = Document(str(file_path))
    texts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            texts.append(text)

    for table in document.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                texts.append(" | ".join(row_texts))

    return "\n".join(texts)


def extract_docx_text_and_images(file_path: Path, ocr_runner) -> list[OCRPage]:
    """Extract native DOCX text and OCR every embedded image."""
    texts: list[str] = []
    native_text = extract_docx_text(file_path)
    if native_text:
        texts.append(native_text)

    with ZipFile(file_path, "r") as archive:
        image_names = sorted(
            name for name in archive.namelist()
            if name.startswith("word/media/")
        )

        for image_name in image_names:
            suffix = Path(image_name).suffix or ".png"
            with NamedTemporaryFile(delete=False, suffix=suffix) as temp:
                temp.write(archive.read(image_name))
                image_path = Path(temp.name)

            try:
                for page in ocr_runner(image_path):
                    if page.text.strip():
                        texts.append(page.text.strip())
            finally:
                image_path.unlink(missing_ok=True)

    return [OCRPage(page=1, text="\n".join(texts), items=[])]
