from enum import Enum
from pathlib import Path
from zipfile import ZipFile

import fitz

from docx import Document


class FileContentType(str, Enum):

    TEXT_ONLY = "text_only"

    IMAGE_ONLY = "image_only"

    TEXT_AND_IMAGE = "text_and_image"

    UNKNOWN = "unknown"


IMAGE_EXTENSIONS = {
    ".jpg",
    ".jpeg",
    ".png",
    ".webp",
    ".bmp",
    ".tif",
    ".tiff",
}


TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".csv",
}


def classify_file(
    file_path: Path,
) -> FileContentType:

    suffix = file_path.suffix.lower()

    # ---------------------------------
    # 이미지 파일
    # ---------------------------------

    if suffix in IMAGE_EXTENSIONS:

        return FileContentType.IMAGE_ONLY

    # ---------------------------------
    # 일반 텍스트 파일
    # ---------------------------------

    if suffix in TEXT_EXTENSIONS:

        return FileContentType.TEXT_ONLY

    # ---------------------------------
    # PDF
    # ---------------------------------

    if suffix == ".pdf":

        return classify_pdf(
            file_path
        )

    # ---------------------------------
    # DOCX
    # ---------------------------------

    if suffix == ".docx":

        return classify_docx(
            file_path
        )

    return FileContentType.UNKNOWN


def classify_pdf(
    file_path: Path,
) -> FileContentType:

    has_text = False
    has_image = False

    document = fitz.open(
        file_path
    )

    try:

        for page in document:

            # -----------------------------
            # 텍스트 확인
            # -----------------------------

            text = page.get_text(
                "text"
            )

            if text and text.strip():

                has_text = True

            # -----------------------------
            # 이미지 확인
            # -----------------------------

            images = page.get_images(
                full=True
            )

            if images:

                has_image = True

            # -----------------------------
            # 둘 다 확인되면 종료
            # -----------------------------

            if (
                has_text
                and has_image
            ):

                return (
                    FileContentType.TEXT_AND_IMAGE
                )

        if has_text:

            return FileContentType.TEXT_ONLY

        if has_image:

            return FileContentType.IMAGE_ONLY

        return FileContentType.UNKNOWN

    finally:

        document.close()


def classify_docx(
    file_path: Path,
) -> FileContentType:

    has_text = False
    has_image = False

    document = Document(
        str(file_path)
    )

    # ---------------------------------
    # 일반 문단
    # ---------------------------------

    for paragraph in document.paragraphs:

        if paragraph.text.strip():

            has_text = True

            break

    # ---------------------------------
    # 표 내부 텍스트
    # ---------------------------------

    if not has_text:

        for table in document.tables:

            for row in table.rows:

                for cell in row.cells:

                    if cell.text.strip():

                        has_text = True

                        break

                if has_text:
                    break

            if has_text:
                break

    # ---------------------------------
    # DOCX 내부 이미지
    # ---------------------------------

    with ZipFile(
        file_path,
        "r",
    ) as archive:

        has_image = any(
            name.startswith(
                "word/media/"
            )
            for name in archive.namelist()
        )

    # ---------------------------------
    # 결과
    # ---------------------------------

    if (
        has_text
        and has_image
    ):

        return (
            FileContentType.TEXT_AND_IMAGE
        )

    if has_text:

        return FileContentType.TEXT_ONLY

    if has_image:

        return FileContentType.IMAGE_ONLY

    return FileContentType.UNKNOWN