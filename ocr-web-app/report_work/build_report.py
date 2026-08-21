import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
ANALYSIS = json.loads((BASE / "analysis.json").read_text(encoding="utf-8"))
OUTPUT = BASE.parent / "LLM_영수증_구조화_성능평가_보고서.docx"

BLUE = "2E74B5"
DARK = "17365D"
MUTED = "667085"
LIGHT = "F2F4F7"
PALE_BLUE = "EAF2F8"
PALE_RED = "FDECEC"
PALE_GREEN = "EAF7EF"
WHITE = "FFFFFF"
BLACK = "111827"


def set_font(run, name="Malgun Gothic", size=10.5, bold=None, color=BLACK, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, headers, rows, widths, aligns=None, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(str(text)), size=font_size, bold=True, color=DARK)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            p = cells[i].paragraphs[0]
            p.alignment = (aligns[i] if aligns else WD_ALIGN_PARAGRAPH.LEFT)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            set_font(p.add_run(str(value)), size=font_size, color=BLACK)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(text), size={1: 16, 2: 13, 3: 12}[level], bold=True,
             color={1: BLUE, 2: BLUE, 3: DARK}[level])
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        set_font(p.add_run(bold_lead), bold=True)
        set_font(p.add_run(text[len(bold_lead):]))
    else:
        set_font(p.add_run(text))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.167
    set_font(p.add_run(text))
    return p


def add_callout(doc, title, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(title), size=11, bold=True, color=DARK)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    set_font(p2.add_run(text), size=10)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def pct(value):
    return f"{value * 100:.1f}%"


def value_text(value):
    if value is None or value == "":
        return "미출력"
    if isinstance(value, (int, float)):
        return f"{value:,.0f}"
    return str(value)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Malgun Gothic"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for level, size, before, after, color in ((1, 16, 16, 8, BLUE), (2, 13, 12, 6, BLUE), (3, 12, 8, 4, DARK)):
    style = styles[f"Heading {level}"]
    style.font.name = "Malgun Gothic"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

# Running header and footer.
hp = section.header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(hp.add_run("OCR 영수증 구조화 | 모델 성능평가"), size=8.5, color=MUTED)
fp = section.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(fp.add_run("내부 평가 자료 · 2026-08-21"), size=8, color=MUTED)

# Memo masthead.
kicker = doc.add_paragraph()
kicker.paragraph_format.space_after = Pt(3)
set_font(kicker.add_run("MODEL PERFORMANCE REVIEW"), size=9, bold=True, color=BLUE)
title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(4)
set_font(title.add_run("LLM 영수증 구조화 성능평가 보고서"), size=23, bold=True, color=DARK)
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(14)
set_font(subtitle.add_run("TEST01~TEST20 일괄평가 결과 및 원본 라벨 재검증"), size=12.5, color=MUTED)

metadata = [
    ("평가 범위", "영수증 이미지 20건 × 모델 2종 = 40회 추론"),
    ("평가 모델", "gemma2:2b / gemma2-yes-category:latest"),
    ("통계 원본", "finance-model-batch-20-statistics (1).json"),
    ("정답 원본", "test01_test20_ground_truth.json"),
    ("보고서 기준일", "2026-08-21"),
]
for label, value in metadata:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(f"{label}: "), size=9.5, bold=True, color=DARK)
    set_font(p.add_run(value), size=9.5)

doc.add_paragraph().paragraph_format.space_after = Pt(2)
add_callout(
    doc,
    "핵심 결론",
    "두 모델 모두 현재 품질 게이트를 통과하지 못했다. gemma2-yes-category는 기본 gemma2보다 품목 검출과 상호 인식이 크게 개선됐지만, "
    "카테고리·총수량·할인액·카드번호·품목 금액에서 운영 적용 수준에 미달한다. 또한 원본 정답의 일부 필드가 평가 요청에서 누락되어 공식 점수는 그대로 최종 의사결정에 사용하기 어렵다.",
    PALE_RED,
)

add_heading(doc, "1. 평가 개요", 1)
add_body(doc, "본 보고서는 다운로드된 일괄 통계 JSON의 공식 점수를 검토하고, 별도 라벨링 JSON 20건을 다시 연결하여 필드별 성능을 재계산한 결과다. 공식 선정 지표(95점 추출 + 속도 3점 + 로컬 비용 2점)와 실제 필드 매칭률을 분리해 해석했다.")
add_bullet(doc, "샘플 수: 영수증 20건, 모델별 성공 응답 20건(호출 성공률 100%)")
add_bullet(doc, "평가 대상: 상호, 날짜, 합계금액, 결제수단, 총수량, 할인액, 카드번호, 카테고리, 품목명·수량·단가·품목금액")
add_bullet(doc, "품질 게이트: JSON 스키마 성공률 98% 이상, 총 결제액 정확도 95% 이상")

add_heading(doc, "2. 공식 선정 지표 결과", 1)
s1 = ANALYSIS["summaries"]["gemma2:2b"]["official"]
s2 = ANALYSIS["summaries"]["gemma2-yes-category:latest"]["official"]
official_rows = [
    ["gemma2:2b", f"{s1['extraction_score_95']:.1f}/95", pct(s1["schema_success_rate"]), pct(s1["total_amount_accuracy"]), f"{s1['average_latency_ms']/1000:.1f}초", f"{s1['final_score_100']:.1f}", "실패"],
    ["gemma2-yes-category", f"{s2['extraction_score_95']:.1f}/95", pct(s2["schema_success_rate"]), pct(s2["total_amount_accuracy"]), f"{s2['average_latency_ms']/1000:.1f}초", f"{s2['final_score_100']:.1f}", "실패"],
]
add_table(doc, ["모델", "추출 점수", "스키마", "합계 정확도", "평균 시간", "최종 점수", "게이트"], official_rows,
          [2100, 1300, 1100, 1250, 1150, 1200, 1260],
          [WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 6, font_size=8.2)
add_body(doc, "gemma2-yes-category는 최종 59.2점으로 gemma2:2b(43.4점)보다 15.8점 높다. 평균 응답시간은 44.2초로 1.3초 느리지만, 정확도 개선 폭에 비해 속도 손실은 작다. 다만 두 모델 모두 스키마와 합계금액 게이트를 충족하지 못한다.")

add_heading(doc, "3. 원본 정답 재검증 결과", 1)
add_callout(
    doc,
    "평가 데이터 무결성 주의",
    "라벨링 원본에는 카테고리, 총 물품 수량, 카드번호가 존재하지만, 다운로드 통계의 각 run.normalized_ground_truth에는 이 필드들이 누락되어 있었다. "
    "특히 프런트엔드 변환은 카테고리를 품목 내부에서 찾도록 되어 있어 원본의 최상위 카테고리를 전달하지 못한다. 아래 '보정 필드 매칭률'은 원본 라벨을 다시 적용한 진단용 수치이며, 공식 100점 점수를 대체하는 재채점은 아니다.",
    PALE_BLUE,
)
c1 = ANALYSIS["summaries"]["gemma2:2b"]
c2 = ANALYSIS["summaries"]["gemma2-yes-category:latest"]
corrected_rows = [
    ["gemma2:2b", pct(c1["stored_mean_field_accuracy"]), pct(c1["corrected_mean_field_accuracy"]), f"{c1['complete_matches']}/20", f"{c1['zero_item_outputs']}/20", pct(c1["item_count_accuracy"])],
    ["gemma2-yes-category", pct(c2["stored_mean_field_accuracy"]), pct(c2["corrected_mean_field_accuracy"]), f"{c2['complete_matches']}/20", f"{c2['zero_item_outputs']}/20", pct(c2["item_count_accuracy"])],
]
add_table(doc, ["모델", "저장 매칭률", "보정 매칭률", "완전일치", "품목 0건 출력", "품목 수 정확도"], corrected_rows,
          [2200, 1400, 1400, 1300, 1500, 1560],
          [WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 5, font_size=8.5)
add_body(doc, "원본 정답을 복원하면 평균 필드 매칭률은 gemma2:2b 27.5%, gemma2-yes-category 46.3%다. 후자는 품목을 전혀 반환하지 않은 사례가 2건으로, 기본 모델의 18건보다 크게 적다. 그러나 20건 중 모든 필드가 맞은 영수증은 두 모델 모두 0건이다.")

add_heading(doc, "4. 필드별 성능", 1)
field_labels = {
    "transaction_date": "날짜", "total_amount": "합계금액", "payment_method": "결제수단",
    "merchant": "상호", "total_quantity": "총수량", "discount_amount": "할인액",
    "card_number": "카드번호", "expense_category": "카테고리", "items.count": "품목 수",
    "items.name": "품목명", "items.quantity": "품목 수량", "items.unit_price": "품목 단가",
    "items.total_amount": "품목 금액",
}
field_order = ["transaction_date", "total_amount", "payment_method", "merchant", "total_quantity", "discount_amount", "card_number", "expense_category", "items.count", "items.name", "items.quantity", "items.unit_price", "items.total_amount"]
field_rows = []
for field in field_order:
    field_rows.append([
        field_labels[field],
        pct(c1["field_accuracy"].get(field, 0)),
        pct(c2["field_accuracy"].get(field, 0)),
        f"{(c2['field_accuracy'].get(field, 0)-c1['field_accuracy'].get(field, 0))*100:+.1f}%p",
    ])
add_table(doc, ["필드", "gemma2:2b", "yes-category", "차이"], field_rows,
          [3300, 1900, 2200, 1960],
          [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER], font_size=8.5)
add_body(doc, "강점은 날짜(두 모델 85%), 결제수단(75%), 합계금액(70~75%)이다. 가장 큰 개선은 품목명(+42.5%p), 품목 수(+45.0%p), 품목 수량(+70.0%p)에서 나타났다. 반면 카테고리는 두 모델 모두 0%, 총수량은 10%, 할인액은 실제 할인 영수증 6건 중 0건 정답이었다.")
add_body(doc, "카드번호의 전체 정확도에는 카드번호가 비어 있는 영수증의 정답 일치가 포함된다. 실제 카드번호가 있는 14건만 보면 gemma2:2b는 3건, yes-category는 0건만 맞아 운영상 카드번호 추출 성능은 매우 낮다.")

add_heading(doc, "5. 주요 오류 패턴", 1)
patterns = [
    ("품목 누락", "기본 gemma2는 20건 중 18건에서 items가 비었다. OCR 평문에서 품목명과 가격의 행·열 관계가 약해지면 소형 모델이 불확실한 품목을 통째로 생략하는 경향이 크다."),
    ("품목 금액 키 불일치", "yes-category 출력에 total_amount 대신 total_price가 반복되어 평가와 Excel 구조에서 품목 금액이 누락됐다. 스키마 강제가 충분하지 않다."),
    ("합계·할인 혼동", "상품합계, 할인 전 금액, 공급가액+부가세와 최종 결제액을 혼동했다. 총액은 비교적 높지만 할인액은 실제 할인 6건 모두 실패했다."),
    ("상호와 시설·OCR 문자열 혼동", "상호 정확도는 기본 모델 10%, yes-category 30%에 그쳤다. 지점명, 시설명, URL, 손상된 OCR 문자열을 상호로 선택하거나 null로 반환했다."),
    ("카테고리 고정값", "두 모델 모두 원본 카테고리와 일치하지 않았다. 현재 출력이 '기타'에 편중되어 카테고리 적용 모델이라는 이름과 실제 결과 사이에 차이가 있다."),
    ("총수량 미출력", "두 모델 모두 총수량 정확도가 10%다. 품목 수량을 일부 읽어도 합산값 또는 영수증 명시 총수량으로 연결하지 못했다."),
]
for title_text, description in patterns:
    add_body(doc, f"{title_text}. {description}", bold_lead=f"{title_text}.")

add_heading(doc, "6. 대표 실패 사례", 1)
case_specs = [
    ("test02.jpg", "BUTTER 단일 품목", "OCR 날짜 오류와 별개로 모델은 품목을 빈 배열로 반환했다. gemma2는 최종 결제액 3,120원 대신 공급가액과 부가세를 더한 3,900원을 선택했고, 승인번호만으로 카드 결제를 추정했다."),
    ("test06.jpg", "청년다방 다품목", "기본 모델은 4개 품목을 모두 누락했다. yes-category는 1개만 남기고 단가도 다른 품목의 값과 결합했다. 총액은 맞았지만 품목 수준 구조화가 실패했다."),
    ("test09.jpg", "할인·복수 품목", "yes-category는 최종 합계 7,000원은 맞췄지만 상호, 결제수단, 총수량, 할인액을 놓쳤고 품목명·단가도 실제 품목과 연결하지 못했다."),
    ("test16.jpg", "COS 할인 영수증", "두 모델 모두 최종 합계 157,600원은 맞췄지만 82,500원 할인액과 총수량 3을 놓쳤다. yes-category는 쇼핑백을 중복 생성하고 의류 품목 두 개를 누락했다."),
]
add_table(doc, ["사례", "유형", "관찰 결과"], case_specs, [1450, 1900, 6010],
          [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], font_size=8.2)

doc.add_page_break()
add_heading(doc, "7. 모델 선정 판단", 1)
add_callout(
    doc,
    "조건부 우선 모델: gemma2-yes-category:latest",
    "두 후보 중에서는 품목 존재 여부와 품목명·수량 재현율이 훨씬 높은 yes-category 모델이 우세하다. 속도 차이는 약 1.3초로 작다. "
    "그러나 현재 점수와 품질 게이트 기준상 바로 운영 모델로 확정할 수준은 아니며, 평가 입력 수정 후 재평가가 선행돼야 한다.",
    PALE_GREEN,
)
add_bullet(doc, "선정 근거: 공식 최종 점수 +15.8점, 보정 필드 매칭률 +18.8%p, 품목 수 정확도 +45.0%p")
add_bullet(doc, "보류 근거: 완전일치 0건, 스키마 66.1%, 합계 정확도 75%, 카테고리 0%, 할인액 유효 사례 0/6")
add_bullet(doc, "기본 gemma2는 품목 0건 출력이 18/20으로 구조화 목적에 부적합")

add_heading(doc, "8. 개선 우선순위", 1)
recommendations = [
    ["P0", "평가 입력 수정", "최상위 카테고리·총수량·카드번호를 전달한 뒤 20건 재평가", "점수 신뢰성 확보"],
    ["P0", "스키마 키 강제", "품목 금액을 total_amount로 고정하고 total_price를 정규화", "품목 금액 누락 감소"],
    ["P0", "후처리 범위 축소", "카드 라벨이 명시된 경우에만 결제수단 보정", "허위 결제수단 감소"],
    ["P1", "품목 최소 검증", "명확한 품목 후보와 빈 items가 충돌하면 재질문·검토 전환", "품목 0건 방지"],
    ["P1", "금액 우선순위", "최종 결제액·할인액·상품합계를 분리하고 산술값은 검산에만 사용", "합계·할인 개선"],
    ["P1", "카테고리 출력", "허용 목록과 OCR 근거를 주고 '기타' 조건 제한", "카테고리 0% 개선"],
    ["P2", "OCR 표 구조 전달", "좌표 기반 행·열 텍스트를 LLM 입력에 추가", "품목 연결 개선"],
]
add_table(doc, ["우선", "과제", "조치", "기대 효과"], recommendations, [750, 1700, 4710, 2200],
          [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT], font_size=8.0)

add_heading(doc, "9. 결론", 1)
add_body(doc, "gemma2-yes-category는 기본 모델보다 우수하지만 운영 확정에는 정확도와 평가 신뢰성이 부족하다. 평가 입력 누락과 품목 금액 키·결제수단·카테고리 문제를 수정한 뒤 동일 20건을 재평가해야 한다.")
add_body(doc, "권고 판정: gemma2-yes-category:latest를 개선·재평가 대상 1순위로 유지하되, 품질 게이트 통과 전 운영 확정은 보류한다.", bold_lead="권고 판정:")

add_heading(doc, "부록. 해석 시 주의사항", 2)
for note in (
    "공식 점수는 다운로드 통계 JSON 값을 그대로 인용했다.",
    "보정 매칭률은 run의 prediction을 원본 라벨과 다시 비교한 진단 지표다.",
    "OCR 원문에 정답이 없는 경우는 모델 단독 책임으로 해석하지 않는다.",
    "선택 필드는 빈 값 일치보다 실제 값이 있는 사례의 성공 건수를 함께 본다.",
):
    p = add_bullet(doc, note)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        set_font(run, size=9.5)

doc.core_properties.title = "LLM 영수증 구조화 성능평가 보고서"
doc.core_properties.subject = "TEST01~TEST20 일괄평가"
doc.core_properties.author = "OCR Web App 성능평가"
doc.core_properties.keywords = "OCR, LLM, 영수증, 구조화, 성능평가"
doc.save(OUTPUT)
print(OUTPUT)
